"""Parse uploaded materials into plain text, and chunk them for retrieval.

Text/markdown work with the stdlib. PDF and DOCX need the `docs` extra
(`pip install -e '.[docs]'`) and are imported lazily so the rest of the tool
runs without them.
"""

from __future__ import annotations

import codecs
from collections.abc import Iterator
from pathlib import Path

# Extracted-text ceiling for compressed formats. PDF/DOCX are zip/deflate
# containers: a small file can decompress to gigabytes ("decompression bomb")
# and OOM the ingest. The caps below bound *actual* decompression by streaming,
# never by trusting a size the file declares about itself (an attacker forges
# that — zipfile still decompresses the real stream).
MAX_EXTRACTED_CHARS = 10_000_000
MAX_DOCX_DECOMPRESSED_BYTES = 200 * 1024 * 1024
MAX_PDF_PAGES = 10_000
MAX_PDF_PAGE_CHARS = 2_000_000  # a single page over this is a bomb signal → truncate
_DECOMP_CHUNK = 1 << 16  # 64 KiB


def _capped_join(parts, sep: str = "\n") -> str:
    """Join text parts, hard-bounded to MAX_EXTRACTED_CHARS.

    Truncates the *current* part to the remaining budget before appending, so a
    single gigantic part can't be materialized into the accumulator whole (the
    old version appended each part first, then capped — unbounded peak)."""
    out: list[str] = []
    budget = MAX_EXTRACTED_CHARS
    first = True
    for part in parts:
        if budget <= 0:
            break
        if not first:
            budget -= len(sep)
        first = False
        piece = part if len(part) <= budget else part[:budget]
        out.append(piece)
        budget -= len(piece)
    return sep.join(out)


def read_document(path: str | Path) -> str:
    """Return the plain-text content of a file (best effort)."""
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(p)
    if suffix == ".docx":
        return _read_docx(p)
    # Default: treat as text, deciding on the CONTENT rather than the suffix — a
    # plain-text material can be a .jsonl export, a .tex draft or an extensionless
    # README, and refusing those would lose real content.
    try:
        raw = p.read_bytes()
    except OSError as exc:
        # Do NOT swallow this. Returning "" makes the caller's `.strip()` filter
        # drop the file with no entry in `skipped`, so a permission-denied file —
        # or one deleted between the scan and the read — vanishes from the corpus
        # silently. parse_materials' per-file handler turns this into a report.
        raise ValueError(f"{p.name}: {exc.strerror or exc}") from exc
    return _decode_text(p.name, raw)


def _decode_text(name: str, raw: bytes) -> str:
    """Decode a material's bytes as text, or raise ValueError if they aren't text.

    Decoding everything with ``errors="replace"`` is what let a .xlsx or a .png
    count as a successfully ingested material: the mojibake is dense
    non-whitespace, so the caller's `.strip()` check keeps it, the file's real
    content never reaches the extractor, and one image can fill the extraction
    window on its own. Refusing here puts the file in ``skipped``, where the
    owner can see it and convert it."""
    # UTF-32 FIRST: its little-endian BOM (ff fe 00 00) starts with the UTF-16
    # one, so testing UTF-16 first decodes a UTF-32 file into NUL-interleaved
    # characters — dense non-whitespace that passes the caller's non-empty check
    # and gets ingested as the garbage this function exists to refuse.
    for bom_pair, encoding in (((codecs.BOM_UTF32_LE, codecs.BOM_UTF32_BE), "utf-32"),
                               ((codecs.BOM_UTF16_LE, codecs.BOM_UTF16_BE), "utf-16")):
        if raw.startswith(bom_pair):
            # Notepad's "Unicode" .txt is one of these — real text, just not
            # UTF-8, which would read it as NUL-interleaved nonsense.
            try:
                text = raw.decode(encoding)
            except UnicodeDecodeError as exc:
                raise ValueError(
                    f"{name}: declares a {encoding.upper()} byte-order mark but could not be "
                    f"decoded as {encoding.upper()} ({exc.reason})") from exc
            if "\x00" in text:  # a BOM is not proof: verify the decode produced text
                raise ValueError(f"{name}: has a {encoding.upper()} byte-order mark but does not "
                                 "decode to text")
            return text
    # Before trying UTF-8, not after: NUL is a perfectly valid UTF-8 character, so
    # a BOM-less UTF-16 file (half NULs) "decodes" cleanly and would be returned
    # as the NUL-interleaved nonsense this function exists to catch.
    if _nul_dense(raw):
        # Either BOM-less UTF-16/32 (iconv -t UTF-16LE, bcp, some CRM exports) or
        # an actual binary file. Only the former decodes to readable text.
        # Score every candidate rather than taking the first that decodes: the
        # WRONG endianness of ASCII-dominant text is valid CJK, which is
        # printable, so a first-match loop reads a UTF-16-BE document as
        # Mandarin. The right endianness of ordinary prose is dominated by
        # low code points; the wrong one is not.
        # -1.0, not 0.0: a genuinely CJK document scores zero on this metric and
        # must still be accepted — the score ranks candidates, it does not gate them.
        best, best_score = None, -1.0
        for encoding in ("utf-16-le", "utf-16-be", "utf-32-le", "utf-32-be"):
            try:
                text = raw.decode(encoding)
            except UnicodeDecodeError:
                continue
            if not _mostly_readable(text):
                continue
            score = sum(1 for ch in text if ord(ch) < 0x250) / len(text)
            if score > best_score:
                best, best_score = text, score
        if best is not None:
            return best
        raise ValueError(f"{name}: looks like a binary file (image, spreadsheet, archive), not text")
    try:
        # Sparse NULs are damage in an otherwise-good document, not a reason to
        # discard it — but they must not travel on into prompts and embeddings.
        return raw.decode("utf-8-sig").replace("\x00", "")   # drops a BOM if present
    except UnicodeDecodeError:
        pass
    # cp1252 maps almost every byte, so it can only be the LAST resort: applied
    # to UTF-8 that is merely damaged (one character truncated mid-sequence, a
    # smart quote spliced into an export) it silently turns every multi-byte
    # character in the whole file into mojibake. Repair the UTF-8 instead when
    # that is what this is.
    if _mostly_utf8(raw):
        return raw.decode("utf-8", errors="replace").replace("\x00", "")
    try:
        return raw.decode("cp1252").replace("\x00", "")
    except UnicodeDecodeError:
        raise ValueError(f"{name}: could not be decoded as text (unrecognized encoding)") from None


def _nul_dense(raw: bytes) -> bool:
    """True if NULs are frequent enough to mean "not a UTF-8 text document".

    A ratio, not mere presence: one stray NUL in an otherwise-valid export is a
    damaged byte, and discarding the whole document over it loses real content —
    with a message that points the owner at the wrong problem. Binary files and
    UTF-16 text are both NUL-dense throughout; the caller tells them apart by
    trying to decode.

    The threshold sits well above what damage produces and well below what an
    encoding produces: UTF-16 of mostly-ASCII text is about half NULs, while a
    document with a few corrupted bytes is a fraction of one percent."""
    if not raw:
        return False
    window = raw[:65536]
    return window.count(0) / len(window) > 0.10


def _mostly_readable(text: str) -> bool:
    """True if a decode produced text rather than a plausible-looking accident.

    Decoding UTF-16 with the wrong endianness rarely fails outright — it yields
    characters from unrelated scripts — so a successful decode is not evidence on
    its own."""
    if not text:
        return False
    ok = sum(1 for ch in text if ch.isprintable() or ch in "\r\n\t")
    return "\x00" not in text and ok / len(text) > 0.9


def _mostly_utf8(raw: bytes) -> bool:
    """True if the bytes are UTF-8 apart from a few damaged spots.

    Distinguishes "UTF-8 with a bad byte" — one character truncated mid-sequence,
    a smart quote spliced into an export — from "not UTF-8 at all", a genuine
    legacy cp1252 document. The test is whether the file carries more multi-byte
    characters that DID decode than bytes that did not: cp1252 text has none that
    decode, while damaged UTF-8 is mostly intact by definition. Getting this
    backwards is expensive in one direction only — cp1252 maps almost every byte,
    so misreading UTF-8 as cp1252 silently mojibakes the whole document."""
    repaired = raw.decode("utf-8", errors="replace")
    bad = repaired.count("�")
    decoded_multibyte = sum(1 for ch in repaired if ord(ch) > 0x7F and ch != "�")
    return decoded_multibyte > bad


def _read_pdf(p: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - optional extra
        raise RuntimeError(
            "PDF parsing needs the `docs` extra:  pip install -e '.[docs]'"
            "  (in your ai-calibrator clone)"
        ) from exc
    reader = PdfReader(str(p))
    # Page-by-page with an immediate per-page + running truncation, so a bomb
    # page can't accumulate into a multi-GB join. (Residual: pypdf materializes
    # one page's text internally before we truncate it — a single crafted page
    # can still spike memory. Bounded per file; see SECURITY.md. Only ingest
    # PDFs you trust.)
    out: list[str] = []
    budget = MAX_EXTRACTED_CHARS
    for i, page in enumerate(reader.pages):
        if i >= MAX_PDF_PAGES or budget <= 0:
            break
        text = page.extract_text() or ""
        if len(text) > MAX_PDF_PAGE_CHARS:
            text = text[:MAX_PDF_PAGE_CHARS]
        if len(text) > budget:
            text = text[:budget]
        out.append(text)
        budget -= len(text)
    return "\n".join(out)


def _zip_decompressed_size_capped(z, cap: int) -> None:
    """Stream-decompress every zip member, aborting past `cap` bytes.

    Reads the real deflate stream in bounded chunks and discards them, so peak
    memory is one chunk — NOT the decompressed size. Trusts nothing the archive
    self-declares: `ZipInfo.file_size` is attacker-controlled and does not bound
    what `zipfile` actually decompresses."""
    total = 0
    for info in z.infolist():
        with z.open(info) as fh:
            while True:
                chunk = fh.read(_DECOMP_CHUNK)
                if not chunk:
                    break
                total += len(chunk)
                if total > cap:
                    raise ValueError(
                        f"decompressed size exceeds the "
                        f"{cap // (1024 * 1024)} MB safety cap (possible zip bomb)"
                    )


def _read_docx(p: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - optional extra
        raise RuntimeError(
            "DOCX parsing needs the `docs` extra:  pip install -e '.[docs]'"
            "  (in your ai-calibrator clone)"
        ) from exc
    import zipfile
    # Bound the REAL decompression first; only then hand a verified-safe file to
    # python-docx (which would otherwise decompress an unbounded stream).
    try:
        with zipfile.ZipFile(p) as z:
            _zip_decompressed_size_capped(z, MAX_DOCX_DECOMPRESSED_BYTES)
    except zipfile.BadZipFile as exc:
        raise ValueError(f"{p.name}: not a valid .docx (zip) file") from exc
    document = docx.Document(str(p))
    return _capped_join(_docx_paragraph_text(document))


def _docx_paragraph_text(container) -> Iterator[str]:
    """Every paragraph's text in `container`, INCLUDING inside its tables.

    `Document.paragraphs` is the body's OWN paragraphs — a table is a separate
    block type, and its cells' paragraphs are not in that list. An FAQ or policy
    matrix (the standard Word layout for one) therefore contributed nothing at
    all, and a table-only file read as "", which ingest counts as a material
    while none of its content reaches the facts, the gaps or the index.

    Walked in document order where python-docx offers it, so a policy row keeps
    the heading it sits under; a cell is a block container itself, so nested
    tables recurse. The fallback covers an older python-docx without
    `iter_inner_content`: tables after prose beats losing them."""
    inner = getattr(container, "iter_inner_content", None)
    items = inner() if inner is not None else [*container.paragraphs, *container.tables]
    for item in items:
        rows = getattr(item, "rows", None)  # a table; a paragraph has none
        if rows is None:
            yield item.text
        else:
            for row in rows:
                for cell in row.cells:
                    yield from _docx_paragraph_text(cell)


def chunk_text(text: str, *, size: int = 1000) -> list[str]:
    """Split text into ~`size`-char chunks by packing whole paragraphs.

    Paragraphs (blank-line separated) stay intact where possible; an oversized
    single paragraph is hard-split as a fallback.

    ``size`` must be a positive integer; ``size <= 0`` is rejected (a zero step
    would otherwise raise an opaque ``range()`` error in the hard-split path).
    """
    if not isinstance(size, int) or size < 1:
        raise ValueError(f"chunk size must be an integer >= 1 (got {size!r})")
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        if current and len(current) + len(para) + 2 > size:
            chunks.append(current)
            current = para
        else:
            current = f"{current}\n\n{para}" if current else para
    if current:
        chunks.append(current)

    # Hard-split any chunk that is still much larger than the target.
    out: list[str] = []
    for chunk in chunks:
        if len(chunk) <= size * 2:
            out.append(chunk)
        else:
            for i in range(0, len(chunk), size):
                out.append(chunk[i : i + size])
    return out
